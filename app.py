import os
import re
import json
import requests
from flask import Flask, request, render_template_string, session
from flask_session import Session
from openai import OpenAI
from uuid import uuid4

client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))

def http_get(url: str) -> str:
    try:
        r = requests.get(url, timeout=5)
        return r.text[:3000]
    except Exception as e:
        return f"HTTP_GET_ERROR: {e}"

def echo(text: str) -> str:
    return f"ECHO_RESULT: {text}"

class Agent:
    def __init__(self, name: str, system_prompt: str, tools=None, model="gpt-4o-mini", memory_file=None):
        self.name = name
        self.system_prompt = system_prompt
        self.tools = tools or {}
        self.model = model
        self.memory_file = memory_file or f"{self.name}_memory.json"
        self.memory = self.load_memory()

    def load_memory(self):
        if os.path.exists(self.memory_file):
            try:
                with open(self.memory_file, "r", encoding="utf-8") as f:
                    return json.load(f)
            except:
                return []
        return []

    def save_memory(self):
        with open(self.memory_file, "w", encoding="utf-8") as f:
            json.dump(self.memory, f, ensure_ascii=False, indent=2)

    def _openai_call(self, messages, functions=None):
        kwargs = {
            "model": self.model,
            "messages": messages,
            "temperature": 0.2,
            "max_tokens": 800
        }
        if functions:
            kwargs["functions"] = functions
            kwargs["function_call"] = "auto"
        resp = client.chat.completions.create(**kwargs)
        return resp.choices[0].message

    def handle(self, user_text: str, session_memory=None) -> str:
        session_memory = session_memory if session_memory is not None else []
        session_memory.append({"role": "user", "content": user_text})
        messages = [{"role": "system", "content": self.system_prompt}] + session_memory
        msg = self._openai_call(messages)
        reply = msg.content.strip()
        session_memory.append({"role": "assistant", "content": reply})
        self.memory.extend(session_memory)
        self.save_memory()
        return reply


# -------- Knowledge helpers ----------
KNOWLEDGE_FILE = "knowledge_base.json"

def load_knowledge():
    if os.path.exists(KNOWLEDGE_FILE):
        try:
            with open(KNOWLEDGE_FILE, "r", encoding="utf-8") as f:
                data = json.load(f)
                return data.get("items", [])
        except:
            return []
    return []

def save_knowledge(items):
    with open(KNOWLEDGE_FILE, "w", encoding="utf-8") as f:
        json.dump({"items": items}, f, ensure_ascii=False, indent=2)

def add_knowledge(text):
    items = load_knowledge()
    items.append(text)
    save_knowledge(items)


# -------- JSON extraction ----------
def extract_json(text: str) -> str:
    if not isinstance(text, str):
        return text
    t = text.strip()
    if t.startswith("```"):
        t_inner = t.strip("`").lstrip()
        if t_inner.lower().startswith("json"):
            t_inner = t_inner[4:].lstrip()
        return t_inner.strip()
    if t.startswith("`") and t.endswith("`"):
        return t.strip("`").strip()
    obj_match = re.search(r"(\{(?:.|\n)*\})", t)
    if obj_match:
        return obj_match.group(1)
    arr_match = re.search(r"(\[(?:.|\n)*\])", t)
    if arr_match:
        return arr_match.group(1)
    return t


# -------- Test plan enrichment ------------

def format_missing_coverage_for_html(item, coverage_summary, missing_coverage_list, rationale_list):
    missing_coverage_text = (
        "With this test case:\n" +
        "\n".join([f"- {c}" for c in coverage_summary]) + "\n\n" +
        "Missing coverage / what to be added:\n" +
        "\n".join([f"- {m}" for m in missing_coverage_list]) + "\n\n" +
        "Rationale of adding / what can be achieved after adding:\n" +
        "\n".join([f"- {r}" for r in rationale_list])
    )
    item["missing_coverage"] = f"<pre>{missing_coverage_text}</pre>"
    item["rationale"] = (
        "<pre>\n" +
        "\n".join([f"- {r}" for r in rationale_list]) +
        "</pre>"
    )
    return item


# -------- NEW: fallback wrapper (minimal patch) ----------
def enforce_formatting_fallback(item):
    # If already formatted correctly, do nothing
    if isinstance(item.get("missing_coverage"), str) and "<pre>" in item["missing_coverage"]:
        return item

    raw_missing = (item.get("missing_coverage") or "").strip()
    raw_rationale = (item.get("rationale") or "").strip()

    fallback_text = (
        "With this test case:\n"
        "- (Original missing_coverage was returned in plain text)\n\n"
        "Missing coverage / what to be added:\n"
        f"- {raw_missing or 'None'}\n\n"
        "Rationale of adding / what can be achieved after adding:\n"
        f"- {raw_rationale or 'Not provided'}"
    )

    item["missing_coverage"] = f"<pre>{fallback_text}</pre>"
    item["rationale"] = f"<pre>- {raw_rationale or 'Not provided'}</pre>"
    return item


# -------- Core enrichment logic ----------
def enrich_test_plan(plan_data):
    for idx, item in enumerate(plan_data.get("plan", [])):
        steps_text = " ".join(item.get("test_case_steps", [])).lower()

        coverage_summary = ["Test steps executed successfully"]
        missing_coverage_list = ["None identified"]
        rationale_list = ["This test plan covers the essential functionalities."]

        # Vcredist classification
        if any(k in steps_text for k in ["vcredist", "visual c++", "vc++", "runtime"]):
            coverage_summary = [
                "Client has vcredist installed",
                "Applications relying on vcredist can run successfully"
            ]
            missing_coverage_list = [
                "Verify vcredist post-installation for all clients",
                "Validate vcredist upgrade paths and old client handling",
                "Deploy applications relying on vcredist and validate"
            ]
            rationale_list = [
                "Ensures the installation process installs required runtime correctly",
                "Ensures upgrades don’t break dependent applications",
                "Confirms clients can run apps dependent on vcredist"
            ]

        # FIRST pass: apply your main formatting
        formatted = format_missing_coverage_for_html(
            item, coverage_summary, missing_coverage_list, rationale_list
        )

        # SECOND pass (only activates if LLM gave broken output)
        plan_data["plan"][idx] = enforce_formatting_fallback(formatted)

    return plan_data




# -------- Flask app ----------
app = Flask(__name__)
app.secret_key = os.getenv("FLASK_SECRET_KEY") or str(uuid4())
app.config["SESSION_TYPE"] = "filesystem"
app.config["SESSION_FILE_DIR"] = "./flask_session_files"
app.config["SESSION_PERMANENT"] = False
Session(app)

UPLOAD_FOLDER = "./uploads"
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

HTML_PAGE = """
<!doctype html>
<html>
<head>
<title>Akili Prime</title>
<style>
    body { 
        font-family: Arial, sans-serif; 
        background: #f5f5f5; 
        margin: 0; 
        padding: 20px;
    }
    .container {
        max-width: 1100px;
        margin: auto;
        background: #fff;
        padding: 20px;
        border-radius: 8px;
        box-shadow: 0 0 10px rgba(0,0,0,0.1);
    }
    h2 img {
        height: 40px;
        vertical-align: middle;
        margin-right: 10px;
    }
    textarea {
        width: 100%;
        font-size: 14px;
        padding: 10px;
    }
    table {
        width: 100%;
        border-collapse: collapse;
        margin-top: 15px;
        background: #fff;
    }
    th, td {
        border: 1px solid #ccc;
        padding: 8px;
        vertical-align: top;
    }
    th { background: #eee; }
    .history-box {
        max-height: 300px;
        overflow-y: auto;
        background: #fafafa;
        padding: 10px;
        border-radius: 6px;
        border: 1px solid #ddd;
        margin-top: 20px;
    }
    .msg-user { color: #1565c0; font-weight: bold; }
    .msg-assistant { color: #2e7d32; font-weight: bold; }
</style>
</head>

<body>
<div class="container">

<h2>
    <img src="/static/agent_icon.png">
</h2>

<form method="post" enctype="multipart/form-data">
    <textarea name="user_input" rows="5" placeholder="Enter your request..."></textarea><br><br>
    <input type="file" name="file"><br><br>
    <input type="submit" value="Send"/>
</form>

{% if table %}
<h3>Prioritized Test Plan</h3>
{{ table|safe }}
{% endif %}

<h3>Conversation History</h3>
<div class="history-box">
{% for entry in history %}
    <p>
        <span class="msg-{{ 'user' if entry.role=='You' else 'assistant' }}">
            {{ entry.role }}:
        </span>
        {{ entry.content|safe }}
    </p>
{% endfor %}
</div>

</div>
</body>
</html>
"""


agent = Agent(
    name="test_optimizer",
    system_prompt=(
        "You are an expert QA assistant. Only generate JSON test plans when requested.\n"
        "JSON structure: { 'plan':[{'functional_area':'...', 'test_case_steps':['...'], 'expected_result':'...', 'missing_coverage':'...'}]}"
    ),
    tools={"http_get": http_get, "echo": echo}
)


@app.route("/", methods=["GET", "POST"])
def chat():
    if "session_memory" not in session:
        session["session_memory"] = []

    chat_history_file = "chat_history.json"
    chat_history = []
    if os.path.exists(chat_history_file):
        with open(chat_history_file, "r", encoding="utf-8") as f:
            chat_history = json.load(f)

    table_html = None

    if request.method == "POST":
        user_input = request.form.get("user_input", "") or ""
        uploaded_file = request.files.get("file")
        url_input = request.form.get("url", "").strip()
        save_k = request.form.get("save_knowledge") == "on"

        knowledge_items = load_knowledge()
        knowledge_block = "\n\n--- STORED KNOWLEDGE ---\n" + "\n\n".join(knowledge_items) if knowledge_items else ""
        transient_sources = []

        # URL content
        if url_input:
            try:
                fetched = http_get(url_input)
                fetched_text = re.sub(r"<[^>]+>", "", fetched)
                transient_sources.append(f"[URL CONTENT FROM {url_input}]\n{fetched_text}")
                if save_k:
                    add_knowledge(f"[URL {url_input}]\n{fetched_text}")
            except Exception as e:
                transient_sources.append(f"[URL CONTENT FROM {url_input} ERROR: {e}]")

        # Uploaded file
        if uploaded_file:
            filename = uploaded_file.filename
            file_path = os.path.join(UPLOAD_FOLDER, filename)
            uploaded_file.save(file_path)
            file_text = ""
            try:
                if filename.lower().endswith(".docx"):
                    from docx import Document
                    doc = Document(file_path)
                    file_text = "\n".join([p.text for p in doc.paragraphs])
                    for table in doc.tables:
                        for row in table.rows:
                            file_text += " | ".join(cell.text for cell in row.cells) + "\n"
                elif filename.lower().endswith(".pdf"):
                    import PyPDF2
                    with open(file_path, "rb") as f:
                        reader = PyPDF2.PdfReader(f)
                        for page in reader.pages:
                            extracted = page.extract_text() or ""
                            file_text += extracted + "\n"
                else:
                    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                        file_text = f.read()
            except Exception as e:
                file_text = f"[UNREADABLE FILE: {e}]"

            transient_sources.append(f"[FILE CONTENT: {filename}]\n{file_text}")
            if save_k:
                add_knowledge(f"[FILE {filename}]\n{file_text}")

        sccm_reference = "You need to generate a complete and detailed test plan.\nReference: https://learn.microsoft.com/en-us"
        combined_parts = [p for p in [knowledge_block] if p]
        if transient_sources:
            combined_parts.append("\n\n--- SUBMISSION SOURCES ---\n" + "\n\n".join(transient_sources))
        combined_parts.append("\n\n--- SCCM REFERENCE ---\n" + sccm_reference)
        combined_parts.append("\n\n--- USER QUERY ---\n" + user_input)
        combined = "\n\n".join(combined_parts)

        chat_history.append({"role": "You", "content": user_input})

        if any(keyword in user_input.lower() for keyword in ["sccm", "test plan"]):
            combined = "Please generate a detailed test plan in JSON format with complete test steps after understanding user requirement and from uploaded documents:\n\n" + combined

        reply = agent.handle(combined, session_memory=session["session_memory"])

        try:
            cleaned = extract_json(reply)
            data = json.loads(cleaned)
            data = enrich_test_plan(data)
            plan = data.get("plan", [])
            rows = ""
            for p in plan:
                steps_list = p.get("test_case_steps", [])
                steps = "<br>".join([f"{i+1}. {s}" for i, s in enumerate(steps_list)])
                rows += (
                    "<tr>"
                    f"<td>{p.get('functional_area', '')}</td>"
                    f"<td>{steps}</td>"
                    f"<td>{p.get('expected_result', '')}</td>"
                    f"<td>{p.get('missing_coverage', '')}</td>"
                    "</tr>"
                )
            if rows:
                table_html = (
                    "<table border='1' style='border-collapse: collapse;'>"
                    "<tr><th>Functional Area</th><th>Test Steps</th><th>Expected Result</th><th>Missing Coverage</th></tr>"
                    f"{rows}</table>"
                )
                chat_history.append({"role": "assistant", "content": table_html})
            else:
                chat_history.append({"role": "assistant", "content": reply})
        except Exception:
            chat_history.append({"role": "assistant", "content": reply})

        with open(chat_history_file, "w", encoding="utf-8") as f:
            json.dump(chat_history, f, ensure_ascii=False, indent=2)

    return render_template_string(HTML_PAGE, history=chat_history, table=table_html)


if __name__ == "__main__":
    os.makedirs(UPLOAD_FOLDER, exist_ok=True)
    os.makedirs("./flask_session_files", exist_ok=True)
    app.run(host="0.0.0.0", port=5000, debug=True)
